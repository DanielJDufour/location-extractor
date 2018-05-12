from sys import version_info
python_version = version_info.major

from broth import Broth
import csv
from datetime import datetime
from docx import Document
from itertools import chain
from os.path import dirname, realpath
from os import listdir
from re import findall, finditer, IGNORECASE, MULTILINE, search, sub, UNICODE
import signal
flags = MULTILINE|UNICODE

try: from language_detector import detect_language
except ImportError: detect_language = None

try: from date_extractor import extract_date
except ImportError: extract_date = None

try: from bscrp import isJavaScript
except ImportError: isJavaScript = None

try: from PyPDF2 import PdfFileReader
except ImportError: PdfFileReader = None


directory_of_this_file = dirname(realpath(__file__))
directory_of_keywords = directory_of_this_file + "/keywords"
directory_of_letters = directory_of_this_file + "/letters"
languages = listdir(directory_of_keywords)

global dictionary_of_keywords
dictionary_of_keywords = {}

global dictionary_of_letters
dictionary_of_letters = {}

global nonlocations
nonlocations = []

open_kwargs = {}
if python_version == 3:
    open_kwargs["encoding"] = "utf-8"

def load_non_locations():
    global nonlocations
    with open(directory_of_this_file + "/nonlocations.txt", **open_kwargs) as f:
        for line in f:
            if line and not line.startswith("#"):
                nonlocations.append(line.strip())

def flatten(lst):
    result = []
    for element in lst: 
        if isinstance(element, (list, tuple)):
            result.extend(flatten(element))
        else:
            for e in element.split(","):
                e = e.strip()
                if e:
                    result.append(e)
    return result

def read_lines(f):
    if python_version == 2:
        return f.read().decode("utf-8").strip().split("\n")
    elif python_version == 3:
        return f.read().strip().split("\n")

def load_language_into_dictionary_of_keywords(language):
    global dictionary_of_keywords
    dictionary_of_keywords[language] = {}
    directory_of_language_files = directory_of_keywords + "/" + language
    language_files = listdir(directory_of_language_files)
    #print "language_files are", language_files
    for filename in language_files:

        # ignore hidden and temporary files like .listed.swp
        if not filename.startswith("."):
            name = filename.split(".")[0]
            #print "filename is", filename
            if filename == "demonyms.txt":
                dictionary_of_keywords[language]['demonyms'] = {}
                with open(directory_of_language_files + "/" + filename, **open_kwargs) as f:
                    for line in read_lines(f):
                        if line:
                            demonym, place = line.split("\t")
                            dictionary_of_keywords[language]['demonyms'][demonym] = place

            elif filename == "abbreviations.txt":
                dictionary_of_keywords[language]['abbreviations'] = {}
                with open(directory_of_language_files + "/" + filename, **open_kwargs) as f:
                    for line in read_lines(f):
                        if line:
                            abbreviation, place = line.split("\t")
                            dictionary_of_keywords[language]['abbreviations'][abbreviation] = place
            elif filename == "countries.txt":
                dictionary_of_keywords[language]['countries'] = {}
                with open(directory_of_language_files + "/" + filename, **open_kwargs) as f:
                    for line in read_lines(f):
                        if line:
                            country, country_code = line.split("\t")
                            dictionary_of_keywords[language]['countries'][country] = country_code
            else:
                with open(directory_of_language_files + "/" + filename, **open_kwargs) as f:
                    keywords = [keyword for keyword in read_lines(f) if keyword]
                    #print "keywords:", keywords
                    keywords += [keyword.title() for keyword in keywords]
                    dictionary_of_keywords[language][name] = keywords

def load_language_into_dictionary_of_letters(language):
    global dictionary_of_letters
    letters = set()
    with open(directory_of_letters + "/" + language + ".txt", **open_kwargs) as f:
        for line in read_lines(f):
            if line:
                letter = line.strip()
                if letter and letter != " ":
                    letters.add(letter)
    dictionary_of_letters[language] = list(letters)


def extract_locations_from_text(text, debug=False, return_demonyms=False, return_abbreviations=False, return_country_info=False, case_insensitive=None):
    global dictionary_of_keywords
    start = datetime.now()
    if debug: print("starting extract_locations_from_text")
    if debug: print("text:", str([text])[:500])
    if debug: print("return_demonyms:", [return_demonyms])
    if debug: print("return_abbreviations:", [return_abbreviations])
    if (python_version == 2 and isinstance(text, str)) or (python_version == 3 and isinstance(text, bytes)):
        text = text.decode("utf-8")

    if case_insensitive is True:
        flags = MULTILINE|UNICODE|IGNORECASE
    else:
        flags = MULTILINE|UNICODE


    # see if we can figure out what language the text is
    # if we can't figure it out, just try to extract locations using all the languages we got
    languages = []

    # language_detector isn't reliable enough yet
    #if detect_language:
    #    detected_language = detect_language(text)
    #    if detected_language:
    #        languages = [detected_language]

    if not languages: 
        languages = listdir(directory_of_keywords)

    if debug: print("languages = ", languages)

    locations = set()
    demonyms = []
    abbreviations = []
    countries = []
    city_country = []
    city_state = []

    for language in languages:

        if debug: print("for language", language)

        if language not in dictionary_of_keywords:
            load_language_into_dictionary_of_keywords(language)

        if language == "English":

            d = dictionary_of_keywords[language]

            #if debug: print "countries for english:", d['countries']


            # \u00e3 is the a with the curly on top like in Sao Tome
            # \u00e9 is the e with accent on top
            location_pattern = br"((?:[A-Z][a-z\u00e3\u00e9]+|[A-Z]|Santo{2,})(?: de)?(?: [A-Z][a-z\u00e3\u00e9]+|[A-Z]{2,})?)".decode('raw_unicode_escape')

            #keyword comes before location
            locations.update(flatten(findall(br"(?:(?:[^A-Za-z\u00e3\u00e9]|^)(?:".decode('raw_unicode_escape') + "|".join(d['before']) + ") )" + location_pattern, text, flags)))
            #print "\nbefore locations are", locations

            #keywords comes after location
            locations.update(flatten(findall(location_pattern + r" (?:" + "|".join(d['after']) + ")", text, flags)))

            #keyword is something that a country can possess, like a president (e.g., "America's President")
            locations.update(flatten(findall(location_pattern + r"'s (?:" + "|".join(d['possessed']) + ")", text, flags)))

            #keyword is something can be listed, like the islands of Santo Domingo and Cuba
            locations_pattern = location_pattern + "((?:, " + location_pattern + ")*),? and " + location_pattern
            pattern = r"(?:" + "|".join(d['listed']) + ")" + r"(?: | of |, |, [a-z]* )" + locations_pattern
            #print "text is", text
            #print "pattern is", pattern
            locations.update(flatten(findall(pattern, text, flags)))

            #ignore demonyms for now, because accuracy is not that high
            #Eritreans, Syrian
            endings = list(set([demonym[-2:] for demonym in list(d['demonyms'].keys())]))
            if debug: print("endings:", endings)
            pattern = r"([A-Z][a-z]{2,}(" + "|".join(endings) + "))"
            if debug: print("pattern:", [pattern])
            for m in finditer(pattern, text, MULTILINE):
                demonym = m.group(0)
                if debug: print("demonym:", demonym)
                if demonym in d['demonyms']:
                    location = d['demonyms'][demonym]
                    demonyms.append({"demonym": demonym, "location": location})

            abbreviation_pattern = r"(?<=[ ,])([A-Z]{2})(?=[ .,]|$)"
            for m in finditer(abbreviation_pattern, text, MULTILINE):
                abbreviation = m.group(0)
                #print "d keys:", d.keys()
                if abbreviation in d['abbreviations']:
                    location = d['abbreviations'][abbreviation]
                    abbreviations.append({"abbreviation": abbreviation, "admin1code": abbreviation, "end": m.end(), "location": location, "start": m.start()})

            for m in finditer(location_pattern + ",? (?P<state>[A-Z]{2})(?=[ .,]|$)", text, MULTILINE):
                name_of_place = m.group(1)
                state = m.group("state")
                city_state.append({"location": name_of_place, "admin1code": state})

                # filter out abbreviation if capture separately because people probably don't want California if have Los Angeles, CA
                start_of_abbreviation = m.start("state")
                end_of_abbreviation = m.end("state")
                abbreviations = [a for a in abbreviations if a["start"] != start_of_abbreviation or a["end"] != end_of_abbreviation]


            # countries
            for m in finditer("(" + "|".join(list(d['countries'].keys())) + ")", text, MULTILINE):
                country = m.group(0)
                if country in d['countries']:
                    countries.append({"country": country, "end": m.end(), "start": m.start(), "country_code": d['countries'][country]})

            if debug: print("countries found are:", countries)

            for m in finditer(location_pattern + ",? ?(?P<country>" + "|".join(d['countries']) + ")", text, MULTILINE):
                if debug: print("found place in country")
                name_of_place = m.group(1)
                country = m.group("country")
                city_country.append({"location": name_of_place, "country": country, "country_code": d['countries'][country]})

                # filter out country if capture separately because people probably don't want England if have London, England
                start_of_country = m.start("country")
                end_of_country = m.end("country")
                countries = [a for a in countries if a["start"] != start_of_country or a["end"] != end_of_country]

            if debug: print("city_country:", city_country)


        elif language == "Arabic":

            if language not in dictionary_of_letters:
                load_language_into_dictionary_of_letters(language)
            letters = dictionary_of_letters[language]
            d = dictionary_of_keywords[language]
            arabic_letter = u"[" + "".join(letters) + "]"
	    #location_pattern = "(" + arabic_letter + br"{3,}(?: \u0627\u0644".decode('raw_unicode_escape') + arabic_letter + "{3,})*)"
            location_pattern = ""

            locations.update(flatten(findall(r"(?:"+ "|".join(d['before']) + ") " + location_pattern, text, flags)))
            locations.update(flatten(findall(location_pattern + r" (?:" + "|".join(d['after']) + ")", text, flags))) # we're using space or underscore bc of tweets

            # could probably speed this up somehow
            # will need to speed this up if list gets longer
            for demonym in d['demonyms']:
                if demonym in text:
                    location = d['demonyms'][demonym]
                    demonyms.append({"demonym": demonym, "location": location})

        elif language == "Spanish":
            d = dictionary_of_keywords[language]
            location_pattern = "((?:(?:[A-Z][a-z]+) )?(?:de )?[A-Z][a-z]+|[A-Z]{2,})"
            locations.update(flatten(findall(r"(?:(?:[^A-Za-z]|^)(?:"+ "|".join(d['before']) + ") )" + location_pattern, text, flags)))
            locations.update(flatten(findall(location_pattern + r" (?:" + "|".join(d['after']) + ")", text, flags)))

            # en una lista
            locations_pattern = location_pattern + "((?:, " + location_pattern + ")*),? y " + location_pattern
            pattern = r"(?:" + "|".join(d['listed']) + ")" + r"(?: | de |, |, [a-z]* )" + locations_pattern
            locations.update(flatten(findall(pattern, text, flags)))

            # demonyms / gentilicos
            for m in finditer(r"[a-z]{3,}(a|e|o|as|es|os)", text, MULTILINE):
                demonym = m.group(0)
                if demonym in d['demonyms']:
                    location = d['demonyms'][demonym]
                    demonyms.append({"demonym": demonym, "location": location})

    # filter out things we often capture that aren't locations
    # and that are actually names of random places
    if not nonlocations:
        load_non_locations()


    locations = [location for location in locations if location not in nonlocations]


    #convert locations to a list
    locations = list(locations)

    #sometimes accidentally picked up a demonymn as a location for a foreign language
    #for example, he is a Libyan will think Libyan is a location because places follow a in Spanish
    list_of_found_demonyms = [d['demonym'] for d in demonyms]
    locations = [location for location in locations if location not in list_of_found_demonyms]
    
    list_of_found_abbreviations = [d['abbreviation'] for d in abbreviations]
    locations = [location for location in locations if location not in list_of_found_abbreviations]


    if return_demonyms or return_abbreviations:
        if return_demonyms:
            locations += demonyms
        if return_abbreviations:
            locations += abbreviations
            # swap out locations for full dict if have info
            # for example remove "Portland" if have {"location": "Portland", "admin1code": "OR"}
            for d in city_state:
                d_location = d['location']
                if d_location in locations:
                    locations.remove(d_location)
                locations.append(d)
        if return_country_info:
            # swap out locations for full dict if have info
            # for example remove "Portland" if have {"location": "Portland", "admin1code": "OR"}
            cities_in_countries = []
            for d in city_country:
                d_location = d['location']
                if d_location in locations:
                    locations.remove(d_location)
                locations.append(d)

    else:
        locations = list(set(locations + [d['location'] for d in demonyms + city_country]))

    #print "extracting locations from text took", (datetime.now() - start).total_seconds(), "seconds"
    if debug: print("finishing extract_locations_from_text with", len(locations), "locations", locations[:5])
    return locations

def extract_location(inpt, debug=False):
    # returns a random location ... in the future make this so return the first location matched.. but implement as separate method entirely
    return extract_locations(inpt, debug=debug)[0]


def extract_locations_with_context_from_tables(tables, debug=False):
    locations = []
    for table in tables:

        # a table is just a list of rows
        if len(table) > 3:
            if debug: print("more than 10 rows in table!")

            # just assume that the header is the first row
            header = table[0]

            #get location column
            date_column_index = None
            location_column_index = None
            admin1_column_index = None
            for column_index, cell in enumerate(header):
                cell = cell.strip().lower()
                if debug: print("cell:", [cell])
                if cell == "city":
                    location_column_index = column_index
                elif cell == "state":
                    admin1_column_index = column_index
                elif "date" in cell or "time" in cell:
                    date_column_index = column_index
            if debug:
                print("location_column_index:", location_column_index)
                print("admin1_column_index:", admin1_column_index)
                print("date_column_index:", date_column_index)

            if location_column_index:
                if debug: print("table:", table)
                for row in table:
                    if debug: print("row:", row)
                    d = {'count': 1}
                    # row is just an array of cells
                    d['name'] = name = row[location_column_index].strip()
                    if admin1_column_index is not None:
                        d['admin1code'] = row[admin1_column_index].strip()
                    if date_column_index is not None:
                        d['date'] = extract_date(row[date_column_index].strip())

                    locations.append(d)

    if debug: print("locations from html table:", locations)
    return locations


def extract_locations_with_context_from_html_tables(text, debug=False):
    try:
        tables = []
        broth = Broth(text)
        if debug: print("broth:", broth)
        for table in broth.tables:
            rows = []
            ths = table.select("thead tr th")
            if ths:
                rows.append([th.text.strip() for th in ths])
            for row in table.select("tr"):
                tds = row.select("td")
                if tds:
                    rows.append([td.text.strip() for td in tds])
            tables.append(rows)
        
        if debug: print("tables:", tables)
        locations = extract_locations_with_context_from_tables(tables, debug=debug)
 
        if debug: print("locations from html table:", locations)
        return locations
    except Exception as e:
        print("[extract_locations_with_context_from_html_tables]:", e)
        raise e

def findnames(text, names, flags=None, debug=True):
    print("names:", type(names))
    result = list(chain.from_iterable([finditer(pattern=name, string=text, flags=flags) for name in names]))
    if debug: print("result:", result)
    return result

def extract_locations_with_context_from_text(text, suggestions=None, ignore_these_names=None, debug=False, max_seconds=None, return_abbreviations=False, case_insensitive=None):
    debug = True
    if debug:
        import inspect
        frame = inspect.currentframe()
        args, _, _, values = inspect.getargvalues(frame)
        print("[location_extractor] starting extract_locations_with_context_from_text with:")
        #for i in args:
            #print("    %s: %s" % (i, [values[i]])[:500])

    start = datetime.now()

    if not extract_date:
        raise Exception("You must have date-extractor installed to use this method.  To fix the problem, run: pip install date-extractor")

    # this is the list of locations that we will return
    locations = []

    # got locations as list of words
    extracted_names = extract_locations_from_text(text, debug=debug, return_demonyms=True, return_abbreviations=return_abbreviations, return_country_info=True, case_insensitive=case_insensitive)
    if debug: print("extracted_names:", extracted_names)
    if ignore_these_names:
        extracted_names = [name for name in extracted_names if name not in ignore_these_names]
        if debug: print("extracted_names after ignoring some:", extracted_names)

    demonym_to_location = {}
    abbreviation_to_location = {}
    name_to_admin1code = {}
    name_to_country = {}
    name_to_country_code = {}
    names = []
    for name in extracted_names:
        if isinstance(name, dict):
            if "demonym" in name:
                demonym_to_location[name['demonym']] = name['location']
                names.append(name['demonym'])
            if "abbreviation" in name:
                abbreviation_to_location[name['abbreviation']] = name['location']
                names.append(name['abbreviation'])
            if "admin1code" in name:
                name_to_admin1code[name['location']] = name['admin1code']
                names.append(name['location'])
            if "country" in name:
                name_to_country[name['location']] = name['country']
                name_to_country_code[name['location']] = name['country_code']
                names.append(name['location'])
        else:
            names.append(name)


    # you can pass in a list of suggested names if you also want to treat those as places
    if suggestions:
        names = set(names + suggestions)


    if debug: print("[location_extractor] finding locations and surrounding information including date and paragraph")
    if debug: counter = -1
    if debug: times_taken = []
    sentences = [{'text': m.group(0), 'start': m.start(), 'end': m.end() } for m in finditer("[^\.\n]+", text)]
    paragraphs = [{'text': m.group(0), 'start': m.start(), 'end': m.end() } for m in finditer("[^\n]+", text)]
    demonym_to_location_keys = list(demonym_to_location.keys())
    if debug: print("[location_extractor] demonym_to_location_keys:", demonym_to_location_keys)
    abbreviation_to_location_keys = list(abbreviation_to_location.keys())
    for matchgroup in findnames(text, names, flags):

        if max_seconds and (datetime.now() - start).total_seconds() > max_seconds:
            break

        if debug:
            start_time_for_mg = datetime.now()
            counter += 1
            if debug: print("matchgroup:", matchgroup)
        name = matchgroup.group(0)
        if debug: print("name:", [name])
        if name:
            middle = float(matchgroup.end() + matchgroup.start()) / 2
            for s in sentences:
                if s['start'] < middle < s['end']:
                    sentence = s['text'].strip()
                    break

            for p in paragraphs:
                if p['start'] < middle < p['end']:
                    paragraph = p['text'].strip()
                    break

            dictionary_of_location = {}

            if any([d in name for d in demonym_to_location_keys]):
                if debug: print("found in demonym_to_location_keys", [demonym_to_location_keys])
                for demonym in demonym_to_location_keys:
                    if debug: "\tdemonym:", [demonym]
                    if demonym in name:
                        if debug: print("\t[location-extractor] changing", name, "to", demonym_to_location[demonym])
                        name = demonym_to_location[demonym]
            elif name in abbreviation_to_location_keys:
                dictionary_of_location['admin1code'] = name
                print("added admin1code to ", dictionary_of_location)
                name = abbreviation_to_location[name]
            elif name in name_to_admin1code:
                dictionary_of_location['admin1code'] = name_to_admin1code[name]
            elif name in name_to_country:
                dictionary_of_location['country'] = name_to_country[name]
                dictionary_of_location['country_code'] = name_to_country_code[name]

            dictionary_of_location['name'] = name

            dictionary_of_location['date'] = date = extract_date(sentence) or extract_date(paragraph)

            if not isJavaScript or not isJavaScript(paragraph):
                dictionary_of_location['context'] = paragraph

            dictionary_of_location['hash'] = str(date) + "-" + name

            locations.append(dictionary_of_location)
            #if debug: print locations
        if debug:
            times_taken.append((datetime.now() - start_time_for_mg).total_seconds())
            print("average_time_taken:", sum(times_taken) / len(times_taken), "seconds")

    if debug: print("[location-extractor] locs before looking in columns:", [l['name'] for l in locations])

    if max_seconds is None or (datetime.now() - start).total_seconds() < max_seconds:
        if debug: print("[location-extractor] get locations by looking for columns with location keywords in them")
        lists = findall("((?:\n^[^\n]{4,20}$){5,})", text, MULTILINE)
        if debug: print("lists:", lists)
        for text_of_list in lists:
            if max_seconds is None or (datetime.now() - start).total_seconds() < max_seconds:
                count = 0
                keywords = []
                for language in dictionary_of_keywords:
                    if "general" in dictionary_of_keywords[language]:
                        keywords += dictionary_of_keywords[language]["general"]
           
                for keyword in keywords: 
                    count += text_of_list.count(keyword)

                if count > 5:
                    lines = text_of_list.split("\n")
                    names = [line for line in lines if line and len(line) > 4 and line not in keywords]

                    for name in names:
                        date = extract_date(text_of_list)
                        locations.append({"name": name, "date": date, "hash": str(date) + "-" + name, "context": name})
    names = [location['name'] for location in locations]


    if debug: print("[location-extractor] locations before removing duplicates:", locations if len(locations) < 5 else len(locations))

    #see: http://stackoverflow.com/questions/21720199/python-remove-any-element-from-a-list-of-strings-that-is-a-substring-of-anothe
    names_verbose = list(filter(lambda x: [x for i in names if x in i and x != i] == [], names))
    if debug: print("[location-extractor] names_verbose:", names_verbose)

    # exclude names_verbose that were passed in suggestions, but not found via regex
    names_verbose = [name for name in names_verbose if (suggestions is None or name not in suggestions) or name in extracted_names]
    if debug: print("[location-extractor] names_verbose:", names_verbose)

    for location in locations:
        name = location['name']
        for name_verbose in names_verbose:
            if not name == name_verbose and name in name_verbose:
                location['name'] = name_verbose

    grouped_by_hash = {}
    if debug: print("[location-extractor] before final loop:", str(locations)[:500], "...")
    for location in locations:
        h = location['hash']
        if h in grouped_by_hash:
            grouped_by_hash[h]['count'] += 1
            #if have same hash, then have same exact location and date, so just need to update context
            if "context" in location and location['context']:
                if 'context' in grouped_by_hash[h]:
                    grouped_by_hash[h]['context'] += location['context']
                else:
                    grouped_by_hash[h]['context'] = "\n ... \n" + location['context']
            if "country" in location and location['country']:
                grouped_by_hash[h]['country'] = location['country']
                grouped_by_hash[h]['country_code'] = location['country_code']
        else:
            d = {'count': 1, 'date': location['date'], 'name': location['name']}
            if "context" in location and location['context']:
                d['context'] = location['context']
            if "admin1code" in location:
                d['admin1code'] = location['admin1code']
            if "country" in location and location['country']:
                d['country'] = location['country']
                d['country_code'] = location['country_code']
            grouped_by_hash[h] = d
 
    locations = list(grouped_by_hash.values())

    #print "extracting locations with context from text took", (datetime.now() - start).total_seconds(), "seconds"
    if debug: print("[location-extractor] finishing extract_locations_with_context_from_text with", list(locations)[:5])
    return locations
       
def extract_locations_from_path_to_pdf(path_to_pdf, debug=False):
    if debug: print("starting extract_locations_from_path_to_pdf", path_to_pdf)
    with open(path_to_pdf, **open_kwargs) as f:
        return extract_locations_from_pdf(pdf)

# takes in a pdf file and returns the text
def get_text_from_pdf_file(pdf_file):
    pdfFileReader = PdfFileReader(pdf_file)
    number_of_pages = pdfFileReader.getNumPages()
    text = ""
    for i in range(number_of_pages):
        text += pdfFileReader.getPage(i).extractText()
    text = sub(" *\n[\n ]*", "\n", text)
    lines = text.split("\n")
    lines = [line for line in lines if len(line) > 4]
    text = "\n".join(lines)
    return text

def get_text_and_tables_from_docx_file(docx_file):
    print("starting get_text_and_tables_from_docx_file")
    document = Document(docx_file)
    print("document:", document)
    text = "\r\n\r\n".join([paragraph.text for paragraph in document.paragraphs])
    print("text:", [text])
    tables = [[[ cell.text for cell in column.cells] for column in table.columns] for table in document.tables]
    print("tables:", tables)
    return {"text": text, "tables": tables}
    
 
def extract_locations_from_pdf(pdf_file):
    #print "starting extract_locations_from_pdf"
    return extract_locations_from_text(get_text_from_pdf_file(pdf_file))

def extract_locations_with_context_from_pdf(pdf_file, debug=False):
    if debug: print("starting extract_locations_with_context_from_pdf with", pdf_file)
    return extract_locations_with_context_from_text(get_text_from_pdf_file(pdf_file))

# will need to make this more sophisticated at a future date
# so that it can pull out tables
def extract_locations_with_context_from_docx(docx_file):
    print("starting extract_locations_with_context_from_docx with", docx_file)
    d = get_text_and_tables_from_docx_file(docx_file)
    locations = []
    locations.extend(extract_locations_with_context_from_text(d['text']))
    locations.extend(extract_locations_with_context_from_tables(d['tables']))
    return locations

def extract_location_with_context(inpt, return_abbreviations=False, debug=False, case_insensitive=None):
    results = extract_locations_with_context(inpt, return_abbreviations=return_abbreviations, debug=debug, case_insensitive=case_insensitive)
    if results:
        return results[0]

def extract_locations(inpt, return_demonyms=False, debug=False):
    if isinstance(inpt, str) or isinstance(inpt, unicode):
        if inpt.endswith(".pdf"):
            return extract_locations_from_path_to_pdf(inpt, debug=debug)
        else:
            return extract_locations_from_text(inpt, return_demonyms=return_demonyms, debug=debug)
    elif isinstance(inpt, file):
        if f.name.endswith(".pdf"):
            return extract_locations_from_pdf(inpt, debug=debug)

def extract_locations_with_context(inpt, suggestions=None, ignore_these_names=[], debug=False, max_seconds=None, return_abbreviations=False, case_insensitive=None):
    if debug: print("starting extract_locations_with_context with", type(inpt))
    print(" [el] type:", str(type(inpt)).lower())
    if isinstance(inpt, str) or (isinstance(inpt, unicode) if python_version == 2 else isinstance(inpt, bytes)):
        if inpt.endswith(".pdf") or inpt.startswith("%PDF"):
            return extract_locations_with_context_from_pdf(inpt, debug=debug)
        else:
            print(" [extract_locations_with_context] calling extract_locations_with_context_from_text")
            return extract_locations_with_context_from_text(inpt, suggestions=suggestions, ignore_these_names=ignore_these_names, debug=debug, max_seconds=max_seconds, return_abbreviations=return_abbreviations, case_insensitive=case_insensitive)
    elif "file" in str(type(inpt)).lower():
        print("isinstance file with name", inpt.name)
        if inpt.name.endswith(".docx"):
            return extract_locations_with_context_from_docx(inpt)
        elif inpt.name.endswith(".pdf"):
            return extract_locations_with_context_from_pdf(inpt)

el = extract_locations
elc = extract_locations_with_context
