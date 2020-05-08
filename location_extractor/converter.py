from docx import Document
from PyPDF2 import PdfFileReader

# takes in a pdf file and returns the text
def get_text_from_pdf_file(pdf_file):
    pdfFileReader = PdfFileReader(pdf_file)
    number_of_pages = pdfFileReader.getNumPages()
    text = ""
    for i in range(number_of_pages):
        text += pdfFileReader.getPage(i).extractText()
    text = sub(" *\n[\n ]*", "\n", text)
    lines = text.split("\n")
    lines = [line for line in lines if len(line) > 4]
    text = "\n".join(lines)
    return text


def get_text_and_tables_from_docx_file(docx_file):
    print("starting get_text_and_tables_from_docx_file")
    document = Document(docx_file)
    print("document:", document)
    text = "\r\n\r\n".join([paragraph.text for paragraph in document.paragraphs])
    print("text:", [text])
    tables = [
        [[cell.text for cell in column.cells] for column in table.columns]
        for table in document.tables
    ]
    print("tables:", tables)
    return {"text": text, "tables": tables}
